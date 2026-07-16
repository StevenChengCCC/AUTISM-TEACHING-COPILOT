from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
from zipfile import BadZipFile, ZipFile

from app.core.config import Settings, settings


_WHITESPACE = re.compile(r"[ \t]+")


class DocumentParsingError(Exception):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    extraction_method: str
    needs_ocr: bool = False
    needs_review: bool = False
    message: str = ""


class V2DocumentParserService:
    def __init__(self, config: Settings = settings):
        self.config = config

    def parse(self, file_name: str, data: bytes) -> ParsedDocument:
        extension = Path(file_name).suffix.lower()
        if extension == ".txt":
            text = self._parse_txt(data)
            method = "txt"
        elif extension == ".pdf":
            text = self._parse_pdf(data)
            method = "pdf_text"
        elif extension == ".docx":
            text = self._parse_docx(data)
            method = "docx"
        else:
            raise DocumentParsingError("This document format is not supported.")
        text = self._normalize(text)
        if extension == ".pdf" and len(text) < self.config.MIN_EXTRACTED_TEXT_CHARS:
            return ParsedDocument(
                text=text,
                extraction_method=method,
                needs_ocr=True,
                message=(
                    "Little or no selectable text was found. This PDF may be scanned; "
                    "OCR or teacher-entered text is required."
                ),
            )
        if not text:
            return ParsedDocument(
                text="",
                extraction_method=method,
                needs_review=True,
                message="No readable text was extracted. Enter or paste the record text.",
            )
        return ParsedDocument(
            text=text,
            extraction_method=method,
            needs_review=True,
            message="Text extracted. Review and correct it before profile extraction.",
        )

    @staticmethod
    def _parse_txt(data: bytes) -> str:
        try:
            return data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentParsingError(
                "TXT files must use UTF-8 text encoding."
            ) from exc

    @staticmethod
    def _parse_pdf(data: bytes) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(data), strict=False)
            return "\n\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception as exc:
            raise DocumentParsingError("The PDF could not be parsed.") from exc

    @staticmethod
    def _parse_docx(data: bytes) -> str:
        try:
            with ZipFile(BytesIO(data)) as archive:
                names = set(archive.namelist())
                if "word/vbaProject.bin" in names:
                    raise DocumentParsingError(
                        "Macro-enabled documents are not allowed."
                    )
                if "word/document.xml" not in names:
                    raise DocumentParsingError(
                        "The DOCX document structure is invalid."
                    )
                if (
                    sum(item.file_size for item in archive.infolist())
                    > 50 * 1024 * 1024
                ):
                    raise DocumentParsingError("The DOCX expanded size is too large.")
            from docx import Document

            document = Document(BytesIO(data))
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except DocumentParsingError:
            raise
        except Exception as exc:
            raise DocumentParsingError("The DOCX file could not be parsed.") from exc

    @staticmethod
    def _normalize(text: str) -> str:
        lines = [_WHITESPACE.sub(" ", line).strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line).strip()
